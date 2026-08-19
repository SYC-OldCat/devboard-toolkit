# -*- coding: utf-8 -*-
"""生成开发板工具箱使用操作文档 (Word 格式)"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "开发板工具箱使用操作手册.docx")

doc = Document()

# ---------- 全局样式 ----------
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.3

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        hs.font.size = Pt(20)
        hs.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        hs.font.size = Pt(15)
        hs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    else:
        hs.font.size = Pt(13)
        hs.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(0.6 * level)
    return p


def add_note(text):
    p = doc.add_paragraph()
    r = p.add_run(f"⚠ {text}")
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    r.font.size = Pt(10)
    return p


def add_tip(text):
    p = doc.add_paragraph()
    r = p.add_run(f"💡 {text}")
    r.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    r.font.size = Pt(10)
    return p


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(cell, '4472C4')
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table


# ====================================================================
# 封面
# ====================================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("开发板工具箱")
r.font.size = Pt(36)
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("使用操作手册")
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run("版本：v2026.08  (最后更新: 2026-08-19)")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ====================================================================
# 目录占位
# ====================================================================
doc.add_heading("目  录", level=1)
toc_items = [
    ("一、工具概览", "3"),
    ("  1.1 功能总览", "3"),
    ("  1.2 文件结构与部署", "3"),
    ("  1.3 配置文件说明 (config.yaml)", "4"),
    ("二、Tab1 数据处理", "6"),
    ("  2.1 输入 / 输出", "6"),
    ("  2.2 ① Jira 数据处理（三种模式）", "6"),
    ("  2.3 ② ADAS 预处理", "9"),
    ("  2.4 执行与取消", "9"),
    ("三、Tab2 感知包编译（Jenkins）", "10"),
    ("四、Tab3 自动回灌", "11"),
    ("  4.1 回灌配置（公共项）", "11"),
    ("  4.2 SDK 回灌子 Tab", "13"),
    ("  4.3 列表回灌子 Tab", "15"),
    ("  4.4 多任务并发（板池共享）", "17"),
    ("  4.5 动态调度 & 增量检测", "18"),
    ("  4.6 失败素材跳过机制", "18"),
    ("  4.7 日志与产物", "19"),
    ("五、Tab4 组合流水线", "20"),
    ("  5.1 节点任选组合", "20"),
    ("  5.2 配置摘要", "20"),
    ("  5.3 一键全流程执行", "21"),
    ("六、常见问题 FAQ", "22"),
]
for t, p in toc_items:
    para = doc.add_paragraph()
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
    para.add_run(t)
    para.add_run(f"\t{p}")
    if not t.startswith("  "):
        for r in para.runs:
            r.bold = True
            r.font.size = Pt(12)

doc.add_page_break()

# ====================================================================
# 一、工具概览
# ====================================================================
doc.add_heading("一、工具概览", level=1)

doc.add_heading("1.1 功能总览", level=2)
doc.add_paragraph("开发板工具箱（Devboard Toolkit）是一套集成化的 ADAS 感知算法验证辅助工具，围绕「素材准备 → 感知包构建 → 开发板回灌」全流程提供自动化能力，主要包含四大功能模块：")

make_table(
    ["Tab", "模块名称", "核心能力"],
    [
        ["Tab1", "数据处理", "Jira 视频拉取 / 视频路径扫描 / 批量复制 + 车型分类 + ADAS 预处理（素材解码、标定对齐）"],
        ["Tab2", "感知包编译", "调用 Jenkins 触发感知包构建任务，自动下载并部署到指定回灌环境"],
        ["Tab3", "自动回灌", "支持 SDK 单目录回灌和列表批量回灌（多板并行、坏素材跳过、动态调度、新板上线利用）"],
        ["Tab4", "组合流水线", "任意组合以上三节点（默认仅自动回灌），并行执行数据处理+编译，串行调用回灌，含总进度条"],
    ],
    col_widths=[2.0, 3.2, 10.5],
)

doc.add_heading("1.2 文件结构与部署", level=2)
doc.add_paragraph("工具采用 PyInstaller 打包为 Windows EXE，部署时请保持以下目录结构：")

p = doc.add_paragraph()
r = p.add_run("部署目录结构：")
r.bold = True

lines = [
    "devboard_toolkit/          ← 部署根目录（可任意命名）",
    "  ├─ devboard_toolkit.exe  ← 主 GUI 程序（双击启动）",
    "  ├─ per_board_runner.exe  ← 单板回灌执行器（被主程序调用，勿手动双击）",
    "  ├─ config.yaml           ← 核心配置文件（开发板、挂载、Jira、Jenkins 等）",
    "  └─ tool/                 ← 工具资源目录",
    "       ├─ fcf_calibration/ ← fcf 标定版本（每个子目录含 vehConfig.json+vruConfig.json）",
    "       │    ├─ default/",
    "       │    ├─ 2021/",
    "       │    └─ 8M_2027/",
    "       └─ lib/             ← 板端运行依赖 .so 库（打包后拷贝到回灌目录）",
]
for line in lines:
    para = doc.add_paragraph()
    run = para.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

add_tip("首次使用前请核对 config.yaml 中的开发板 IP、挂载账号密码、Jenkins 地址、Jira 账号是否正确。")

doc.add_heading("1.3 配置文件说明 (config.yaml)", level=2)
doc.add_paragraph("config.yaml 是工具的全局配置文件，所有 Tab 均依赖此配置。以下分模块说明关键配置项：")

doc.add_heading("1.3.1 boards 开发板配置", level=3)
make_table(
    ["配置项", "说明", "示例"],
    [
        ["boards.<板名>.host", "开发板 SSH IP", "172.17.188.189"],
        ["boards.<板名>.port", "SSH 端口", "22"],
        ["boards.<板名>.user", "SSH 用户名", "root"],
        ["boards.<板名>.password", "SSH 密码（固定 arcsoft123）", "arcsoft123"],
        ["boards.<板名>.timeout", "SSH 连接超时（秒）", "8"],
    ],
    col_widths=[4.5, 6.5, 5.0],
)
add_tip("普通板：board1~board6（共6块）；线上板：Online1~Online7（共7块，需在 GUI 勾选「使用线上开发板」才参与检测/使用）。支持最多13块板并行。")

doc.add_heading("1.3.2 mount 共享挂载配置", level=3)
make_table(
    ["配置项", "说明", "示例"],
    [
        ["mount.source", "UNC 共享路径（素材/回灌根）", "//172.17.12.118/Model_Test/..."],
        ["mount.point", "板端 Linux 挂载点", "/mnt"],
        ["mount.username / password", "SMB 认证凭据", "syc53636"],
        ["mount.testbed_subpath", "testbed 相对共享根的子路径", "/SYC/testbed"],
    ],
    col_widths=[4.5, 6.5, 5.0],
)

doc.add_heading("1.3.3 usage_check 空闲板检测阈值", level=3)
make_table(
    ["配置项", "说明", "默认值"],
    [
        ["loadavg_threshold", "CPU 负载阈值（超过视为使用中）", "4.0"],
        ["net_rx_threshold_gb", "网卡收包量阈值（GB，超过视为使用中）", "1.0"],
    ],
    col_widths=[4.5, 7.5, 4.0],
)

doc.add_heading("1.3.4 replay_env 回灌环境扫描", level=3)
make_table(
    ["配置项", "说明"],
    [
        ["windows_host", "Windows 侧 UNC 主机名替代（如 hz-iotfs02）"],
        ["testbed_subpath", "同 mount.testbed_subpath，用于拼接回灌环境完整路径"],
    ],
    col_widths=[4.0, 11.7],
)

doc.add_heading("1.3.5 car_models 车型-标定映射", level=3)
doc.add_paragraph("键为车型名（GUI 下拉列表的显示值），值为回灌环境中 calibration json 的文件名（不含 .json 后缀）。例：")
p = doc.add_paragraph()
run = p.add_run("  '0452': calibration_result_2M_Fitting_Geely_0452")
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading("1.3.6 replay_list_template / replay_sdk_template", level=3)
doc.add_paragraph("板端 shell 脚本模板（YAML | 字面量块，保留换行）。内含：")
add_bullet("坏素材智能跳过循环（自动识别 testbed 返回码，提取坏素材，写入 failed_*.txt）")
add_bullet("运行清单强制覆盖复制（防止 replay_tmp 残留空文件导致假成功）")
add_bullet("空 running_ 文件直接报错退出（退出码 2）")
add_bullet("LD_LIBRARY_PATH、素材挂载、输出目录等环境变量导出")
add_note("请勿随意修改模板语法；修改后请确保 YAML 格式有效，尤其是 | 字面量块缩进。")

doc.add_heading("1.3.7 jenkins 感知包编译", level=3)
make_table(
    ["配置项", "说明", "示例"],
    [
        ["jenkins.server", "Jenkins HTTP 地址", "http://172.17.189.18:8080"],
        ["jenkins.username / password", "登录凭据", "qatest"],
        ["jenkins.download_dir", "构建产物下载本地目录", r"\\hz-iotfs02\...\pkgs"],
        ["jenkins.default_job", "默认构建任务名", "25640-2_PDT_..._Lq560v200"],
    ],
    col_widths=[4.5, 6.5, 5.0],
)

doc.add_heading("1.3.8 jira_data Jira 数据处理", level=3)
make_table(
    ["配置项", "说明"],
    [
        ["base_url", "Jira 基础 URL（默认 https://jira.arcsoft.com.cn:8443）"],
        ["username / password", "Jira 登录凭据"],
        ["max_workers", "并发下载线程数（上限 8）"],
    ],
    col_widths=[4.0, 11.7],
)

doc.add_heading("1.3.9 adas 预处理", level=3)
make_table(
    ["配置项", "说明"],
    [
        ["exe_path", "ADAS_Visualization.exe 完整路径（未配置则跳过预处理且不报错）"],
        ["timeout", "单素材预处理超时（秒）"],
        ["max_workers", "并发预处理进程数"],
    ],
    col_widths=[4.0, 11.7],
)

doc.add_page_break()

# ====================================================================
# 二、Tab1 数据处理
# ====================================================================
doc.add_heading("二、Tab1 数据处理", level=1)
doc.add_paragraph("数据处理 Tab 包含两个可独立/组合使用的子模块：① Jira 数据处理（素材拉取 + 分类 + 去重）、② ADAS 预处理（素材解码生成 bin/mp4 和 log）。")

doc.add_heading("2.1 输入 / 输出", level=2)
make_table(
    ["字段", "说明"],
    [
        ["输入", "模式不同含义不同：Jira链接/批量复制 = txt 文件路径；视频路径 = 文件夹路径。点击右侧「浏览…」选择。"],
        ["输出目录", "处理后文件落地的根目录。所有模式统一在此目录下生成子文件夹。"],
    ],
    col_widths=[3.0, 12.7],
)

doc.add_heading("2.2 ① Jira 数据处理（三种模式）", level=2)
p = doc.add_paragraph()
r = p.add_run("启用方式：")
r.bold = True
doc.add_paragraph("勾选「启用 Jira 数据处理」复选框。取消勾选则子模块不运行（但参数保留）。")

doc.add_heading("2.2.1 模式一：Jira 链接", level=3)
add_bullet("输入：txt 文件，每行一条 Jira issue 链接（如 https://jira.arcsoft.com.cn:8443/browse/FTIM-1719）或纯 issue_id（如 FTIM-1719，自动补前缀）。")
add_bullet("流程：")
steps = [
    "解析每行 issue，调用 Jira REST API 拉取附件列表（视频 .h265/.h264 等）",
    "下载附件到 输出目录",
    "可选「创建 Jira 子目录」：每个 issue 建独立文件夹（输出目录/FTIM-1719/）",
    "可选「车型分类」：按文件名关键词归类到车型子目录",
    "可选「创建同名文件夹」：每个视频文件创建同名父文件夹包含该视频",
    "可选「只保留最大后缀」：同前缀多数字后缀的文件，仅保留数字最大的版本",
    "（文件名先去除 _cameraId_N 后缀，再比较数字后缀）",
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

doc.add_heading("2.2.2 模式二：视频路径", level=3)
add_bullet("输入：文件夹路径。递归遍历该目录下所有 .h265/.h264 视频。")
add_bullet("流程：")
for i, s in enumerate([
    "递归扫描输入目录下的所有视频文件",
    "跳过已存在的同名输出文件（不覆盖，不计入失败）",
    "复制到输出目录（可选「车型分类」→「创建同名文件夹」）",
    "支持 UNC 路径输入，如 \\\\dtc-fs04\\SmartCar_Collect\\...",
], 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)
add_tip("此模式不显示「创建 Jira 子目录」和「只保留最大后缀」（选项自动隐藏）。")

doc.add_heading("2.2.3 模式三：批量复制", level=3)
add_bullet("输入：txt 文件，每行一个 .h265/.h264 文件的绝对路径（可本地/UNC）。")
add_bullet("行为基本同「视频路径」模式，区别在于输入来源是 txt 清单而非目录扫描。")
add_bullet("支持跳过已存在文件。")

doc.add_heading("2.2.4 通用选项说明", level=3)
make_table(
    ["选项", "默认", "说明"],
    [
        ["并发数", "5", "下载/复制并行线程数，上限 8。越大越快但网络/IO 压力更高。"],
        ["车型分类", "勾选", "按视频名关键词（如 natie3、0452、5013…）匹配车型，创建对应子目录。匹配不上的归入 other/。"],
        ["创建同名文件夹", "勾选", "每个视频 xxx.h265 生成 xxx/xxx.h265（符合回灌目录结构要求）。"],
        ["创建 Jira 子目录", "不勾选", "仅 Jira 链接模式可见。按 issue_id 再套一层子目录。"],
        ["只保留最大后缀", "勾选", "仅 Jira 链接模式可见。如 foo_1.h265 和 foo_3.h265 同时存在，仅保留 foo_3.h265。"],
    ],
    col_widths=[3.5, 1.8, 10.4],
)

doc.add_heading("2.3 ② ADAS 预处理", level=2)
p = doc.add_paragraph()
r = p.add_run("启用方式：")
r.bold = True
doc.add_paragraph("勾选「启用 ADAS 预处理」。注意 Jira 数据处理和 ADAS 预处理可任意组合：都勾选 → 先处理素材再预处理；只勾选前者 → 仅拉取复制；只勾选后者 → 对输出目录现有素材做预处理。")

make_table(
    ["字段", "默认", "说明"],
    [
        ["确认车型", "3 - 其他/商用车", "传 -v 参数给 ADAS_Visualization.exe：0=gl8，1=拿铁，2=理想one，3=其他/商用车，4=吉利ss21，5=五菱f510s，6=旧版商用，7=东湖真值，8=其他3.0协议，9=日产p20n。"],
        ["生成mcap", "否", "传 -m True/False。mcap 是 ROS2 格式产物，按实际需要开启。"],
    ],
    col_widths=[3.0, 3.5, 9.2],
)

add_note("预处理的前提：config.yaml 中 adas.exe_path 正确指向 ADAS_Visualization.exe。若路径不存在，该步骤自动跳过不报错。")

doc.add_heading("2.4 执行与取消", level=2)
make_table(
    ["按钮", "说明"],
    [
        ["▶ 开始执行", "校验必填项（至少启用一个子模块 + 输入 + 输出目录），启动后台线程执行。"],
        ["× 取消", "发送 stop_event，任务将在下一次循环迭代（如下载下一个素材前）中止。已完成文件保留。"],
        ["📂 打开输出目录", "调用系统资源管理器打开输出目录。"],
    ],
    col_widths=[4.0, 11.7],
)
add_tip("日志输出在本 Tab 底部的「执行日志」面板。各 Tab 日志相互独立（使用线程局部 stdout 代理），不会串台。")

doc.add_page_break()

# ====================================================================
# 三、Tab2 感知包编译
# ====================================================================
doc.add_heading("三、Tab2 感知包编译（Jenkins）", level=1)
doc.add_paragraph("本 Tab 通过调用 Jenkins API 触发感知包构建任务，并将构建产物（感知包 + runtime）自动解压到指定的回灌环境目录。")

make_table(
    ["字段", "说明"],
    [
        ["SDK zip", "必选。选择从 Jenkins 下载页或其他来源获取的 SDK_SOURCE_*.zip 文件。内含源码及构建参数。"],
        ["输出目录", "可选。构建产物解压落地的目录（建议填 Tab3 中选的回灌环境完整 UNC 路径）。留空则下载解压到 SDK zip 同目录。"],
    ],
    col_widths=[3.0, 12.7],
)

make_table(
    ["按钮", "说明"],
    [
        ["▶ 编译 & 下载", "调用 jenkins_build.auto_build_main：触发 Jenkins job → 轮询构建状态 → 下载产物 zip → 解压 → 扫描感知包名。"],
        ["× 取消", "停止客户端轮询。注意：服务端的 Jenkins 构建任务会继续运行直至完成。"],
        ["📂 打开输出目录", "打开已填的输出目录（未填则提示）。"],
    ],
    col_widths=[4.0, 11.7],
)

add_tip("成功后日志末尾会打印「感知包名: NH_ADAS_PERCEPTION_xxx」，此名称会被 Tab3 / Tab4 自动识别为候选感知包。")
add_note("Jenkins 配置（server/账号/default_job）见 1.3.7 节。构建超时较长（5~20 分钟），期间可切换其他 Tab 并行操作（如 Tab1 处理素材）。")

doc.add_page_break()

# ====================================================================
# 四、Tab3 自动回灌
# ====================================================================
doc.add_heading("四、Tab3 自动回灌", level=1)
doc.add_paragraph("本 Tab 是工具的核心，负责生成板端脚本 + SSH 连接开发板 + 并行回灌。支持 SDK 单目录回灌 和 列表批量回灌 两种子模式，且支持同一 Tab 内启动多个任务共享板池。")

doc.add_heading("4.1 回灌配置（公共项）", level=2)
make_table(
    ["字段 / 控件", "说明", "操作"],
    [
        ["回灌环境", "UNC 共享下 testbed/ 的子目录（即本次回灌使用的 runtime+感知包 所在目录）。",
         "下拉选择；首次启动自动扫描；或点 🔄 刷新 手动重扫。宽度 56 个字符以显示全称。"],
        ["感知包", "所选回灌环境下匹配 NH_ADAS_PERCEPTION_* 的子目录名。",
         "切换回灌环境后自动刷新。"],
        ["fcf标定文件", "tool/fcf_calibration/ 下的版本名（default / 2021 / fcf_20260812 / 8M_2027 等）。",
         "**直接覆盖策略**：每次回灌都会把选中版本的 vehConfig.json + vruConfig.json 复制覆盖到回灌目录，不校验是否已存在。"],
        ["车型/标定", "config.yaml 中 car_models 配置的键。选中后映射为板端脚本的 {{CALIBRATION}}.json。",
         "下拉选择。宽度 56 个字符。"],
        ["开发板数量", "本任务期望最多使用的板子数量。",
         "Spinbox 1~6（不勾选线上板）或 1~13（勾选线上板）。实际可用数以 🔍 检测空闲板 结果为准。"],
        ["🔍 检测空闲板", "并行（ThreadPoolExecutor，每板独立线程）SSH 登录所有（或已过滤的）板子，检查 CPU loadavg 和网卡收包量，判定空闲/使用中。",
         "检测结果立即流式打印（先完成先打印）。空闲板加入共享板池。"],
        ["使用线上开发板", "是否包含 Online1~Online7 共 7 块线上板。",
         "默认不勾选 → 检测池 6 块板；勾选 → 检测池 13 块板。"],
    ],
    col_widths=[3.2, 7.5, 5.0],
)

add_tip("空闲板判定规则（来自 usage_check.py）：loadavg < usage_check.loadavg_threshold 且 net_rx < net_rx_threshold_gb 才视为空闲。阈值可在 config.yaml 调整。")

doc.add_heading("4.2 SDK 回灌子 Tab", level=2)
doc.add_paragraph("SDK 回灌针对单个素材目录（input/{{USER}}/{{相对路径}}/ 下的视频），传统单路径回灌模式。")

make_table(
    ["字段", "默认", "说明"],
    [
        ["用户名", "空", "即 {{USER}}，用于拼板端 input 路径与 output 路径。建议填英文名首字母缩写。"],
        ["日期", "今天 YYYYMMDD", "自动填充，拼到 output 目录中：output/{{USER}}/{{DATE}}_{{感知包后缀}}/{{车型}}/。"],
        ["素材相对路径", "空", "例：20260810/0452。板端实际素材路径 = $PWD/input/{{USER}}/{{素材相对路径}}/。"],
        ["视频路径（可选）", "空", "新增功能：填一个含 .h265/.h264 的目录后，点击启动会**智能检测每个素材是否已预处理**。"],
    ],
    col_widths=[3.2, 3.5, 8.5],
)

doc.add_heading("4.2.1 视频路径 + 智能预处理检测", level=3)
p = doc.add_paragraph()
r = p.add_run("检测逻辑：")
r.bold = True
doc.add_paragraph("对视频路径下每个 .h265/.h264 文件，执行以下判定：")
for i, s in enumerate([
    "文件本身不为空（size > 0）",
    "同目录下存在同名的 .bin 和 .mp4 文件（两者都非空）",
    "同目录下 log/ 文件夹存在且不为空（里面至少 1 个 .txt 文件）",
], 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

make_table(
    ["判定结果", "处理方式"],
    [
        ["已预处理（上述 3 条全满足）",
         "调用 classify_by_car.py 按车型关键词生成 txt → 启动**列表回灌**流程（共享 Tab3 列表回灌的多板调度、坏素材跳过、增量检测等全部策略）。生成的 txt 命名为 {USER}_{car}.txt，按条目数升序依次处理，跳过 error_unknown.txt。自动根据文件名匹配车型→标定（如 GZY_5013.txt → 选 5013 标定），匹配不到的 txt 跳过并记日志。"],
        ["未预处理（任意一条不满足）",
         "调用 Tab1 数据处理的视频路径逻辑：复制素材到 「素材相对路径」下，创建同名文件夹；再调用 ADAS 预处理（-v 车型 -m mcap）；预处理完成后走**常规 SDK 回灌**（单目录）。混合场景串行顺序：先未预处理、再已预处理。"],
    ],
    col_widths=[3.5, 12.2],
)

add_note("视频路径字段是**可选增强**。不填则 SDK 回灌按原有逻辑直接以素材相对路径回灌。")

doc.add_heading("4.3 列表回灌子 Tab", level=2)
doc.add_paragraph("列表回灌是大规模批量回灌的推荐模式：输入素材清单 txt（或视频路径自动生成 txt），由工具按素材大小 LPT 分片 + 多板并行 + 动态调度接力 + 增量检测新板 + 坏素材自动跳过。")

make_table(
    ["字段 / 控件", "默认", "说明"],
    [
        ["用户名", "空", "同 SDK 回灌。"],
        ["日期", "今天 YYYYMMDD", "同 SDK 回灌。"],
        ["素材输入: txt 文件 ●", "默认选中", "选择一个素材清单 txt（每行一个板端素材绝对路径，如 /tmp/iot_test/mnt_data/SYC/testbed/.../xxx.h265）。浏览默认打开当前选中的回灌环境 UNC 路径。"],
        ["素材输入: 视频路径 ○", "未选", "选择一个含 .h265/.h264 的文件夹；内部调用 classify_by_car.py 生成 {USER}_{car}.txt 清单（同 SDK 视频路径），然后按条目数升序依次处理。"],
    ],
    col_widths=[3.2, 2.5, 9.5],
)

doc.add_heading("4.3.1 分片策略（LPT 按文件大小调度）", level=3)
doc.add_paragraph("为解决「素材大小不均导致最后只剩一块板跑大文件」的问题，列表回灌采用以下分片流程：")
for i, s in enumerate([
    "读取输入 txt 每一行 → UNC→Linux 路径映射（支持多 UNC 前缀共享同一挂载点，迭代查找）→ 并行 ThreadPoolExecutor(32) 获取素材文件 size；PermissionError/访问不到记为 0。",
    "自适应计算每份目标条数：per_target = clamp(ceil(总条数 / (板数 × 3)), 1, 60)。板数使用「配置最大板数」（不勾选线上=6，勾选=13），从而保证后续新板上线也有分片可接。",
    "总份数 N = ceil(总条数 / per_target)，确保每份 ≤ 60 且通常 ≥ 板数×3。",
    "按 size 降序排列所有素材（堆排序等价物），使用最小堆轮询分配到 N 个分片（LPT 算法），保证各分片的「总 size 和」接近一致。乱序对回灌结果无影响。",
    "若所有素材 size 均为 0（访问不到/路径错），回退为顺序分片以保留原始顺序。",
], 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

doc.add_heading("4.3.2 动态调度接力 + 增量检测", level=3)
doc.add_paragraph("列表回灌的主调度循环：")
add_bullet("检测空闲板 → 取 N 份分片 → 先分配给最先空闲的板；板完成当前分片后自动取下一份。")
add_bullet("每次新 txt 启动前 / 同 txt 接力空窗时：重新检测空闲板（包括板池中原本未被占用的板）。per_target 下限 = 1，最大化分板运行。")
add_bullet("增量检测（BOARD_RESCAN_INTERVAL = 60 秒，即 1 分钟一次）：")
subs = [
    "每 1 分钟扫描「当前未被任何任务占用」的板子，发现新空闲板立即加入共享空闲池，后续分配接力使用。",
    "同时反向检查：空闲池内的板如果在新一轮检测中被判定变忙（被其他用户手动占用），则从空闲池移除，防止「幽灵空闲板」误分配。",
]
for s in subs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(s)
    p.paragraph_format.left_indent = Cm(0.8)

add_bullet("提前启动阈值（避免空等）：")
subs2 = [
    "若当前无任何在跑任务 → 阈值=1：只要池中有 ≥1 块板空闲就立即启动下一份分片/下一个 txt。",
    "若当前有在跑任务且 per_target ≤ 15（小分片） → 阈值=1。",
    "若当前有在跑任务且 per_target > 15（正常分片） → 阈值=2：池中有 ≥2 块板空闲时才启动（让更多板凑齐一次并行接力效率更高）。",
]
for s in subs2:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(s)
    p.paragraph_format.left_indent = Cm(0.8)

add_bullet("Reboot 策略：**每个分片启动前（含同 txt 接力）无条件 reboot 板子**，防止板上挂载残留/手动改路径导致回灌失败。每次约 +60 秒。")

doc.add_heading("4.4 多任务并发（板池共享）", level=2)
doc.add_paragraph("Tab3 支持在同一个自动回灌 Tab 内**同时启动多个回灌任务**（例如：一个任务用 3 块板回灌版本 A，另一个任务用 3 块板回灌版本 B）。核心机制：")

for i, s in enumerate([
    "共享板池（_available_pool） + 共享占用集合（_busy_boards） + 锁（_pool_lock）：所有任务的取板/还板操作都串行化经过池。",
    "每个任务维护独立 thread / stop_event / 日志上下文（线程局部 stdout 代理，多任务日志不串台）。",
    "每个任务取板上限 = 用户在「开发板数量」Spinbox 填的数值，而不是池内总板数。所以如果有 6 块空闲，任务 A 选 3、任务 B 选 3，可以并行。",
    "增量检测（1 分钟/次）跨任务共享：检测到的新空闲板，由先请求的任务先分配到。",
], 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

add_note("任务注册表中每个任务有唯一 task_id。点击「× 停止」会发送当前任务的 stop_event，不影响同 Tab 下其他仍在运行的任务。")

doc.add_heading("4.5 失败素材跳过机制", level=2)
doc.add_paragraph("板端 shell 脚本（replay_list_template）内实现 while 循环，让回灌遇到坏素材时**自动跳过并继续下一个，不会中断整个分片**。逻辑概述：")

for i, s in enumerate([
    "先 cp 原始分片到 replay_tmp/running_<分片>.txt（强制覆盖，防止上次残留的空 running_ 文件导致假成功）。然后检查 running_ 非空，为空直接 exit 2。",
    "只要 running_ 不为空就循环：",
    "  (a) 设置 materialInputPath=running_，执行 ./${app_path} 2>&1 | tee tmp_run_$$.log。",
    "  (b) 成功判定：grep 到 testbed->run ret = 0 且未 grep 到失败模式 → 本分片 OK，break。",
    "  (c) 失败判定命中以下任一即视为有坏素材：",
    "      - testbed->run ret = [1-9][0-9]*（非零返回）",
    "      - testbed_error_code:[1-9][0-9]*",
    "      - ERROR ARC_ADAS_Initial ret = [1-9]*（初始化失败，比如示例中的 33685585）",
    "      - No file found in input path / Invalid file path in txt",
    "  (d) 提取坏素材路径：优先用 log 中最后一条 'Processing file: <path>'；匹配不到就兜底取 running_ 首行。",
    "  (e) 坏素材写入 replay_tmp/failed_<分片>.txt；打印 [SKIP] 坏素材: <path> 到终端。",
    "  (f) 从 running_ 中去掉坏素材行及其之前行（awk flag 法）。如果去不掉（行数没变），兜底删除首行并写入 failed。",
    "  (g) 删除临时运行日志，继续下一轮。",
    "循环结束后，若 failed_* 非空打印其路径和条数；否则打印「本分片无失败素材」。",
], 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

add_bullet("死锁防护：MAX_SKIP = 50 次；连续同素材命中即强制跳出，防止脚本死循环。")
add_bullet("每个任务的每个 txt 启动前会自动清 replay_tmp 目录（running_* / failed_* / tmp_run_* / tmp_next_*），避免跨任务污染。")
add_tip("GUI 侧每块板回灌完成后，通过 {板名}.done 标记文件（含板名、exit_code、finish_at）聚合判定所有板完成，然后汇总失败统计。")

doc.add_heading("4.6 日志与产物", level=2)

doc.add_heading("4.6.1 板端 / UNC 侧输出路径", level=3)
make_table(
    ["产物", "路径规则", "说明"],
    [
        ["算法输出", "output/{{USER}}/{{DATE}}_{{APP_SUFFIX}}/{{车型}}/", "由板端脚本设置 outputMountPath。"],
        ["终端日志", "{回灌UNC路径}/logs/{板名}_{感知包后缀}_{时间戳}.log",
         "每个板子一次运行一份。stdout/stderr 同时打印到 Windows Terminal + 写入该 log 文件（行缓冲 buffering=1，立即 flush）。"],
        ["失败素材清单", "{回灌板端目录}/replay_tmp/failed_<分片>.txt",
         "仅当本分片有坏素材时才生成。没有失败则不创建空文件。"],
        ["运行进度清单", "replay_tmp/running_<分片>.txt", "坏素材移除后剩余的待跑列表，用于接力恢复。"],
        ["回灌完成标记", "logs/{板名}.done", "包含 board/exit_code/finish_at，GUI 轮询等待用。"],
    ],
    col_widths=[3.0, 6.0, 6.7],
)

doc.add_heading("4.6.2 GUI 日志格式要点", level=3)
add_bullet("每块板启动分片时打印：「boardX 开始回灌 yaq_463_4.txt」，便于追踪进度。")
add_bullet("板-分片分配会在启动 txt 时集中打印分配明细。")
add_bullet("待回灌聚合显示：「待回灌: X 个新txt + Y 份接力」/「待回灌: X 个新txt」/「待回灌: Y 份接力」/「待回灌: 无」，不会重复。")
add_bullet("增量检测提示：「开发板检测 (N 块, 并行检测中...)」 + 每块板实时流式打印结果。")

make_table(
    ["按钮", "说明"],
    [
        ["▶ 生成脚本 & 启动回灌", "校验参数 → 生成 shell 脚本 → 写入 UNC → 启动 per_board_runner.exe 多板并行 → 轮询等待所有板完成。"],
        ["× 停止", "向当前任务发送 stop_event（多任务模式下不影响其他任务），释放所有占用板回池。"],
        ["📂 打开回灌目录", "打开选中回灌环境的 UNC 目录。"],
        ["回灌结束自动删除脚本", "默认不勾选。勾选则所有板完成后删除生成的 start_*.sh 脚本，保持目录整洁。"],
    ],
    col_widths=[4.5, 11.2],
)

doc.add_page_break()

# ====================================================================
# 五、Tab4 组合流水线
# ====================================================================
doc.add_heading("五、Tab4 组合流水线", level=1)
doc.add_paragraph("Tab4 提供「一键全流程」能力：任选组合 Tab1~Tab3 的节点，先并行执行「数据处理 + 感知包编译」（两者都勾选时），完成后执行自动回灌。")

doc.add_heading("5.1 节点任选组合", level=2)
doc.add_paragraph("流水线节点复选框默认状态：")
make_table(
    ["节点", "默认勾选", "对应 Tab 逻辑"],
    [
        ["① 数据处理", "否", "Tab1：Jira 数据处理（按模式）+ 可选 ADAS 预处理。参数值直接读取 Tab1 已填写/选择的控件。"],
        ["② 感知包编译", "否", "Tab2：Jenkins 构建 + 下载 + 解压。参数值直接读取 Tab2 控件。"],
        ["③ 自动回灌", "是", "Tab3：SDK/列表回灌。参数值直接读取 Tab3 控件（环境、感知包、fcf、车型、板数、子Tab字段等）。"],
    ],
    col_widths=[3.0, 2.0, 10.7],
)
add_tip("至少勾选一个节点才能启动。节点未勾选会自动跳过，并在摘要中显示「(节点未勾选 → 跳过)」。")

doc.add_heading("5.2 配置摘要", level=2)
doc.add_paragraph("点击右上角「🔄 刷新配置摘要」按钮（或启动后自动读取），会把三个 Tab 的参数汇总显示，方便启动前核对：")
add_bullet("① 数据处理摘要：模式（Jira链接/视频路径/批量复制）、车型分类、同名文件夹、Jira子目录、只保留最大后缀、ADAS预处理、并发数。")
add_bullet("② 感知包编译摘要：Jenkins Job 名、SDK zip 文件名、输出目录。节点未勾选时提示使用现有 runtime+感知包。")
add_bullet("③ 自动回灌摘要：方式（SDK/列表 + 列表时是 txt 还是视频路径）、回灌环境名、车型、用户名、感知包、SDK 模式时额外显示素材相对路径。")

doc.add_heading("5.3 一键全流程执行", level=2)

doc.add_heading("5.3.1 Step 0：参数校验", level=3)
doc.add_paragraph("启动后首先做所有勾选节点的必填项校验，任何一项不通过立即终止并打印错误，不产生任何副作用：")
add_bullet("① 勾选 → 输入非空且有效（Jira/batch=文件存在，video=目录存在；输出目录非空）。")
add_bullet("② 勾选 → SDK zip 存在；输出目录非空。")
add_bullet("③ 勾选 → 回灌环境已选、车型已选、感知包已选、fcf 版本索引合法；SDK 模式需用户名+素材相对路径；列表模式需用户名+日期+素材输入存在；子 Tab 字段按模式区分。")

doc.add_heading("5.3.2 Step 1：并行（数据处理 + 感知包编译）", level=3)
add_bullet("仅当至少勾选 ① 或 ② 时执行。两个都勾选 → ThreadPoolExecutor(2) 并行跑，分别有独立的日志上下文。")
add_bullet("任一节点失败或用户点击「× 中止」→ 立即终止整个流水线（失败策略已删除，不再提供「继续/停止/重试 3 次」等选项）。")

doc.add_heading("5.3.3 Step 2：串行自动回灌", level=3)
add_bullet("仅当勾选 ③ 时执行。Step1 全部成功（或跳过）后才开始。")
add_bullet("直接调用 TabFeedback._do_start_task（与 Tab3 手动启动完全相同的代码路径），所以多板调度、坏素材跳过、增量检测、日志等能力完全一致。")

doc.add_heading("5.3.4 控件说明", level=3)
make_table(
    ["控件", "说明"],
    [
        ["▶ 一键全流程执行", "Step0 校验 → Step1 并行（如果有）→ Step2 回灌（如果有）。含独立线程 + stop_event。"],
        ["× 中止", "设置流水线 stop_event。Step1 的两个节点各自检查 stop_event 后退出；Step2 回灌的 stop_event 同样被触发。"],
        ["总进度 Progressbar", "按勾选节点数量等分。例：只勾选③ → 0%→100% 一段；①+③ → 0%→50%（Step1完）→100%（Step2完）。右侧有等待启动/StepX/已完成/已取消/失败等状态文本。"],
        ["流水线日志面板", "独立日志面板（线程局部 stdout），三步骤的日志按时间线聚合显示，不会和 Tab1~Tab3 的其他日志串台。"],
    ],
    col_widths=[4.5, 11.2],
)

doc.add_page_break()

# ====================================================================
# 六、常见问题 FAQ
# ====================================================================
doc.add_heading("六、常见问题 FAQ", level=1)

faqs = [
    ("Q1. 双击 devboard_toolkit.exe 无反应 / 闪退？",
     "① 确认 exe 目录下是否有 config.yaml 和 tool/ 子目录，缺一不可。\n"
     "② 打开 cmd 后运行 exe，看控制台输出的 Traceback。常见原因：config.yaml YAML 格式错误（如 replay_list_template 缩进错）。"),

    ("Q2. 日志字体模糊？",
     "工具已启用 Per-Monitor V2 DPI Awareness，若仍模糊请确认：① 显示器缩放不是「125%（推荐）」以外的特殊值；② 系统显示设置中「修复应用缩放」已开启。"),

    ("Q3. 开发板连接超时 / SSH 认证失败？",
     "① 确认本机可 ping 通板 IP（config.yaml 中 boards.<name>.host）。\n"
     "② 密码固定为 arcsoft123，勿修改。\n"
     "③ 板端 reboot 后约 60~90s 才能恢复 SSH，属正常现象。"),

    ("Q4. 列表回灌有素材没灌上 / 漏掉？",
     "常见原因：① 输入 txt 中有路径 UNC 和板端路径映射不上，导致 size=0 时被顺序分片后，板子上实际文件不存在而跳过。请确认素材路径在板端 /tmp/iot_test/mnt_data 下挂载正确。\n"
     "② 如果看到「本分片无失败素材」但 testbed 没跑：检查 replay_tmp 是否被清理——工具启动时会自动清理；如果手动拷脚本请先删 replay_tmp 目录。"),

    ("Q5. 明明板子都忙，GUI 还显示池中有 N 块空闲板？",
     "这是「幽灵空闲板」。已修复：增量检测时不仅扫描未用的板子，还会反向检测空闲池内的板是否变忙，变忙则移除。请确保使用最新 exe（2026.08 及之后）。"),

    ("Q6. 回灌报错 素材初始化失败（ERROR ARC_ADAS_Initial ret = 33685585）但整个分片直接停止？",
     "老版本模板是一次失败就退出。新版 exe 已启用坏素材跳过循环。确认 config.yaml 中 replay_list_template 是 | 字面量块并且包含 while [ -s \"${_run_list}\" ] 的完整结构。"),

    ("Q7. 待回灌显示「0 个」但其实还有接力分片？",
     "旧版本有格式重复/统计口径不一致的问题。新版日志格式已统一为「待回灌: X 个新txt + Y 份接力」等四种场景之一。请使用最新 exe。"),

    ("Q8. 感知包编译（Jenkins）点了取消，但 Jenkins 网页上还在跑？",
     "正常。取消只停客户端轮询线程，服务端构建不受影响。需要手动去 Jenkins 网页 Abort。"),

    ("Q9. 自动回灌 Tab 启动第二个任务显示「没有空闲板」但实际有？",
     "确认 Spinbox「开发板数量」不要超过池里的板数。多任务模式下每个任务的取板上限就是这个值，所以总板数够但每个任务都取满会被拒绝。先点击 🔍 检测空闲板 刷新池。"),

    ("Q10. 视频路径模式下生成的多个 txt 顺序处理？是否并行？",
     "按 txt 条目数**升序**依次处理，小的先跑完释放板给大的。处理每个 txt 内部是多板并行接力。并且每个新 txt 启动前都会**重新检测空闲板**。同时当前运行的板如果先完成会自动接力跑同 txt 的下一份分片，所以板资源不浪费。"),

    ("Q11. fcf 标定文件是不是只在回灌环境缺失时才复制？",
     "不是（旧行为）。当前是**直接覆盖**策略，不论回灌目录下有没有旧的 vehConfig.json / vruConfig.json，都会复制选中版本覆盖写入。"),

    ("Q12. exe 发其他用户用，需要拷贝哪些文件？",
     "完整拷贝部署目录结构（见 1.2 节）：devboard_toolkit.exe、per_board_runner.exe、config.yaml、tool/ 目录整体。其他用户的 config.yaml 中如果 Jira/Jenkins 账号不同，需要单独改。"),

    ("Q13. 如何确认回灌真的成功了？",
     "三重确认：① 板端日志 grep 'testbed->run ret = 0' 或 'testbed task end' 且无失败模式；② 对应 {板名}.done 标记文件 exit_code=0；③ GUI 日志末尾打印「所有板回灌完成」+「终端日志已保存: xxx.log」。"),
]

for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run(q)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    for line in a.split("\n"):
        pp = doc.add_paragraph(line)
        pp.paragraph_format.left_indent = Cm(0.4)
    doc.add_paragraph()

# ====================================================================
# 保存
# ====================================================================
doc.save(OUTPUT_PATH)
print(f"[✓] 文档已生成: {OUTPUT_PATH}")
print(f"    大小: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")
